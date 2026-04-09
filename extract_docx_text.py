import argparse
import os
import sys


def _iter_paragraph_text(doc):
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t:
            yield t


def _iter_table_text(doc):
    for table in doc.tables:
        for row in table.rows:
            cells = []
            for cell in row.cells:
                text = " ".join((cell.text or "").split())
                cells.append(text)
            # Keep table structure readable in txt
            line = "\t".join(cells).strip()
            if line:
                yield line


def extract_docx_to_text(input_path: str) -> str:
    try:
        from docx import Document
    except Exception as e:  # pragma: no cover
        raise RuntimeError(
            "Missing dependency 'python-docx'. Install it with: python -m pip install python-docx"
        ) from e

    doc = Document(input_path)
    lines = []
    lines.extend(_iter_paragraph_text(doc))

    table_lines = list(_iter_table_text(doc))
    if table_lines:
        if lines:
            lines.append("")
        lines.append("[Tables]")
        lines.extend(table_lines)

    return "\n".join(lines).strip() + "\n"


def main() -> int:
    ap = argparse.ArgumentParser(description="Extract .docx text to a UTF-8 .txt file.")
    ap.add_argument("input", help="Path to .docx file")
    ap.add_argument(
        "-o",
        "--output",
        help="Path to output .txt (defaults to <input>.txt)",
        default=None,
    )
    args = ap.parse_args()

    input_path = os.path.abspath(args.input)
    if not os.path.exists(input_path):
        print(f"Input not found: {input_path}", file=sys.stderr)
        return 2

    out_path = os.path.abspath(args.output) if args.output else (input_path + ".txt")

    try:
        text = extract_docx_to_text(input_path)
    except Exception as e:
        print(str(e), file=sys.stderr)
        return 1

    os.makedirs(os.path.dirname(out_path) or ".", exist_ok=True)
    with open(out_path, "w", encoding="utf-8") as f:
        f.write(text)

    print(out_path)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())

