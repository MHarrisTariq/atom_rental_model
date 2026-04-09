"""
Generate a single documentation bundle: Word (.docx) + Markdown (.md) containing
every project source file in full, with per-file explanations and an architecture update.

Includes: backend Python, frontend source (no node_modules/dist), root scripts, requirements.
Excludes: node_modules, dist, __pycache__, .git, .vite, binaries (.joblib, .zip), dataset CSV.

Run from project root:
  python build_complete_project_documentation.py

Requires: pip install python-docx
"""
from __future__ import annotations

import os
from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.enum.text import WD_LINE_SPACING
from docx.shared import Pt, RGBColor

ROOT = Path(__file__).resolve().parent

OUT_DOCX = ROOT / "Smart_Pricing_Complete_Code_And_Explanations.docx"
OUT_MD = ROOT / "Smart_Pricing_Complete_Code_And_Explanations.md"

SKIP_DIRS = {"node_modules", "dist", "__pycache__", ".git", ".vite"}
SKIP_FILES = {".DS_Store"}
SKIP_SUFFIXES = {".pyc", ".pyo", ".joblib", ".zip", ".csv"}
SKIP_EXACT_NAMES = {"Smart_Pricing_Complete_Code_And_Explanations.md"}  # avoid self-embed if re-run

# Preferred reading order; remaining files are appended sorted.
ORDERED: list[str] = [
    "requirements.txt",
    "requirements_demo.txt",
    "HOW_TO_RUN_SMART_PRICING.txt",
    "run_demo.py",
    "build_client_zip.py",
    "build_full_code_docx.py",
    "extract_docx_text.py",
    "train_lodging_price_model.py",
    "build_complete_project_documentation.py",
    "backend/__init__.py",
    "backend/pricing_engine.py",
    "backend/storage.py",
    "backend/auth.py",
    "backend/rate_limit.py",
    "backend/ml_conversion.py",
    "backend/train_conversion_model.py",
    "backend/decision_engine.py",
    "backend/observability.py",
    "backend/events.py",
    "backend/feature_store.py",
    "backend/main.py",
    "frontend/package.json",
    "frontend/package-lock.json",
    "frontend/tsconfig.json",
    "frontend/vite.config.ts",
    "frontend/index.html",
    "frontend/src/vite-env.d.ts",
    "frontend/src/api.ts",
    "frontend/src/main.tsx",
    "frontend/src/styles.css",
    "frontend/src/App.tsx",
]

EXPLAIN: dict[str, str] = {
    "requirements.txt": (
        "Python dependencies for the upgraded stack: FastAPI, Uvicorn, pandas, scikit-learn, "
        "JWT auth (python-jose), rate limiting (slowapi), Prometheus metrics, and python-docx for tooling."
    ),
    "requirements_demo.txt": (
        "Minimal dependency list referenced by HOW_TO_RUN and older zip handoffs (FastAPI, Uvicorn, pandas). "
        "Prefer requirements.txt for full features."
    ),
    "HOW_TO_RUN_SMART_PRICING.txt": (
        "Human-readable runbook: one-command production flow (build UI, run_demo.py), dev mode with Vite proxy, "
        "and data requirement (Airbnb_Open_Data.csv in project root)."
    ),
    "run_demo.py": (
        "Entry point to run Uvicorn with backend.main:app from project root; serves built React from "
        "frontend/dist when SMART_PRICING_WEB_DIST is set or default path exists."
    ),
    "build_client_zip.py": (
        "Packages backend/, frontend source, CSV, and helper files into SmartPricing_Lodging_Client.zip for delivery, "
        "excluding node_modules, dist, and __pycache__."
    ),
    "build_full_code_docx.py": (
        "Earlier doc generator (subset ordering). Superseded by this script for full coverage; kept for history."
    ),
    "extract_docx_text.py": (
        "CLI utility: extracts paragraph and table text from a .docx to UTF-8 .txt (uses python-docx)."
    ),
    "train_lodging_price_model.py": (
        "Standalone training script for lodging price modeling (if present in repo); may predate conversion model."
    ),
    "build_complete_project_documentation.py": (
        "This file: collects all source paths, writes Word + Markdown with explanations and full file bodies."
    ),
    "backend/__init__.py": (
        "Marks backend as a Python package so `import backend.main` and `uvicorn backend.main:app` work."
    ),
    "backend/pricing_engine.py": (
        "Core deterministic pricing: multiplier pipeline (season, demand, supply, DOW, lead time, quality, goal, risk), "
        "clamp to min/max, smoothing, blackout/locked days, kill-switch flattening, confidence and explanation tags, "
        "plus booking_probability_mock and expected_revenue helpers for simulation fallback."
    ),
    "backend/storage.py": (
        "SQLite persistence (WAL): app_state (kill switch, regional override), host_prefs JSON per listing_id, "
        "append-only audit_log. Env SMART_PRICING_DB_PATH overrides default backend/app.db."
    ),
    "backend/auth.py": (
        "JWT helpers: HS256 access tokens with role claim; FastAPI dependencies get_current_role and require_admin; "
        "secret/algorithm from SMART_PRICING_JWT_SECRET and SMART_PRICING_JWT_ALG."
    ),
    "backend/rate_limit.py": (
        "SlowAPI limiter keyed by client IP; default 120/min; used on API routes; exception handler wired in main."
    ),
    "backend/ml_conversion.py": (
        "Loads joblib artifact (model + feature_order); predict_proba(dict) for conversion probability. "
        "Path from SMART_PRICING_CONVERSION_MODEL_PATH or default backend/models/conversion_model.joblib."
    ),
    "backend/train_conversion_model.py": (
        "Trains a scikit-learn LogisticRegression pipeline on a proxy label derived from CSV fields; "
        "writes conversion_model.joblib. Not ground-truth bookings—placeholder for real labeled data."
    ),
    "backend/decision_engine.py": (
        "Decision layer: builds feature dict from ListingSignals + price; grid search to maximize price × conversion "
        "when ML model is loaded; returns action/price/expected_revenue metadata for simulation API."
    ),
    "backend/observability.py": (
        "Logging level from SMART_PRICING_LOG_LEVEL; request middleware adds x-request-id and x-response-time-ms; "
        "optional Prometheus /metrics via prometheus-fastapi-instrumentator when installed."
    ),
    "backend/events.py": (
        "Pluggable event publisher: no-op by default; KafkaPublisher when SMART_PRICING_KAFKA_BROKERS is set "
        "(kafka-python). main.py publishes pricing.viewed, host.settings_saved, simulation.ran, admin.kill_switch."
    ),
    "backend/feature_store.py": (
        "Abstraction for listing features: in-memory stub or RedisFeatureStore when SMART_PRICING_REDIS_HOST (and port) "
        "set—scaffold for future online features."
    ),
    "backend/main.py": (
        "FastAPI app: CORS, rate limits, JWT login POST /api/auth/login, listings, pricing calendar, host settings "
        "(admin JWT), simulation with ML or mock probability + decision_engine block, admin status/kill switch, "
        "SQLite-backed state, optional Kafka events, Prometheus, static mount for frontend/dist."
    ),
    "frontend/package.json": (
        "NPM metadata: React 18, Vite 5, TypeScript; scripts for dev/build/preview."
    ),
    "frontend/package-lock.json": (
        "Exact transitive dependency versions for reproducible npm installs."
    ),
    "frontend/tsconfig.json": (
        "TypeScript: strict, ES2022, JSX react-jsx, bundler resolution for Vite."
    ),
    "frontend/vite.config.ts": (
        "Vite + React plugin; dev server port 5173; proxy /api to 127.0.0.1:8000."
    ),
    "frontend/index.html": (
        "HTML shell: title, viewport, font link, root div for React mount."
    ),
    "frontend/src/vite-env.d.ts": (
        "Vite client types reference for import.meta and env."
    ),
    "frontend/src/api.ts": (
        "Fetch wrappers for /api/listings, pricing, host settings, simulation, admin. Note: backend may require "
        "Bearer token for admin/settings endpoints—extend headers when using secured APIs."
    ),
    "frontend/src/main.tsx": (
        "React 18 createRoot, StrictMode, App + global CSS import."
    ),
    "frontend/src/styles.css": (
        "Global layout, theme variables, calendar grid, forms, modal, responsive tweaks for Smart Pricing UI."
    ),
    "frontend/src/App.tsx": (
        "Main UI: tabs Host / Simulation / Admin; listing picker; host min/max/base/goals; 60-day calendar with modal; "
        "simulation results; admin kill switch and audit (API calls must align with backend auth when enabled)."
    ),
}


def should_skip_file(rel: str, path: Path) -> bool:
    if rel.replace("\\", "/") in SKIP_EXACT_NAMES:
        return True
    low = rel.lower()
    for token in ("/node_modules/", "/dist/", "/.vite/"):
        if token in low.replace("\\", "/"):
            return True
    if path.suffix.lower() in SKIP_SUFFIXES:
        return True
    return False


def collect_files() -> list[Path]:
    found: dict[str, Path] = {}

    def walk_tree(base: Path, under: str) -> None:
        if not base.exists():
            return
        for dirpath, dirnames, filenames in os.walk(base):
            dirnames[:] = [d for d in dirnames if d not in SKIP_DIRS]
            for fn in filenames:
                if fn in SKIP_FILES:
                    continue
                p = Path(dirpath) / fn
                rel = p.relative_to(ROOT).as_posix()
                if should_skip_file(rel, p):
                    continue
                # Only text-like sources for backend tree
                if under == "backend":
                    if p.suffix.lower() not in {".py"}:
                        continue
                if under == "frontend":
                    rel_norm = rel.replace("\\", "/")
                    if not (
                        rel_norm.startswith("frontend/src/")
                        or rel_norm
                        in {
                            "frontend/package.json",
                            "frontend/package-lock.json",
                            "frontend/tsconfig.json",
                            "frontend/tsconfig.node.json",
                            "frontend/vite.config.ts",
                            "frontend/index.html",
                        }
                    ):
                        continue
                found[rel] = p

    walk_tree(ROOT / "backend", "backend")
    walk_tree(ROOT / "frontend", "frontend")

    # Root scripts and configs
    for name in (
        "requirements.txt",
        "requirements_demo.txt",
        "HOW_TO_RUN_SMART_PRICING.txt",
        "run_demo.py",
        "build_client_zip.py",
        "build_full_code_docx.py",
        "extract_docx_text.py",
        "train_lodging_price_model.py",
        "build_complete_project_documentation.py",
    ):
        p = ROOT / name
        if p.is_file():
            rel = p.relative_to(ROOT).as_posix()
            if not should_skip_file(rel, p):
                found[rel] = p

    ordered_paths: list[Path] = []
    for rel in ORDERED:
        if rel in found:
            ordered_paths.append(found.pop(rel))
    for rel in sorted(found.keys()):
        ordered_paths.append(found[rel])

    return ordered_paths


def add_para(doc: Document, text: str, *, bold: bool = False, size: int = 11) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    if bold:
        run.bold = True
    run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)


def add_code_block_docx(doc: Document, content: str) -> None:
    if not content.endswith("\n"):
        content += "\n"
    for line in content.splitlines():
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        run = p.add_run(line)
        run.font.name = "Consolas"
        run._element.rPr.rFonts.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}ascii", "Consolas")
        run._element.rPr.rFonts.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}hAnsi", "Consolas")
        run.font.size = Pt(7.5)


def write_markdown(paths: list[Path], intro: str) -> None:
    lines: list[str] = []
    lines.append("# Smart Pricing — Complete source code and explanations\n")
    lines.append(f"_Generated: {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')}_\n\n")
    lines.append(intro)
    lines.append("\n---\n\n")

    for fp in paths:
        rel = fp.relative_to(ROOT).as_posix()
        expl = EXPLAIN.get(rel, f"Project source file: `{rel}`.")
        lines.append(f"## `{rel}`\n\n")
        lines.append(f"{expl}\n\n")
        ext = fp.suffix.lower()
        lang = {".py": "python", ".ts": "typescript", ".tsx": "tsx", ".css": "css", ".json": "json", ".html": "html", ".txt": "text"}.get(ext, "")
        body = fp.read_text(encoding="utf-8", errors="replace")
        fence = lang or "text"
        lines.append(f"```{fence}\n")
        lines.append(body)
        if not body.endswith("\n"):
            lines.append("\n")
        lines.append("```\n\n")
        lines.append(f"<!-- end of file: {rel} -->\n\n")

    OUT_MD.write_text("".join(lines), encoding="utf-8")


def main() -> None:
    paths = collect_files()

    intro = """## Scope

This document embeds **every hand-written project source file** for the Smart Pricing (Lodging) demo: **backend (Python/FastAPI)** and **frontend (React/TypeScript/Vite)**, plus root **requirements** and **utility scripts**.

### Excluded (not application source)

- `node_modules/`, `frontend/dist/`, `__pycache__/`, `.vite/`, `.git/`
- Binaries and datasets: `*.joblib`, `*.zip`, `*.csv`
- Third-party vendored code inside `node_modules` (thousands of files; reinstall via `package-lock.json`)

### Architecture update (enterprise-oriented)

1. **Persistence**: SQLite (`backend/storage.py`) stores kill-switch, regional override, host settings, and audit log (env `SMART_PRICING_DB_PATH`).
2. **Security**: JWT login (`POST /api/auth/login`); admin-only routes use `require_admin` (`backend/auth.py`). Rate limiting via `slowapi` (`backend/rate_limit.py`).
3. **ML**: Trained conversion model (`backend/train_conversion_model.py`) saved as `backend/models/conversion_model.joblib`; inference in `backend/ml_conversion.py`. Simulation uses ML when the file exists, else mock curve.
4. **Decision engine**: `backend/decision_engine.py` grid-searches price to maximize expected revenue using the conversion model.
5. **Observability**: Request IDs + optional `/metrics` (`backend/observability.py`).
6. **Events / features (optional)**: Kafka publisher (`backend/events.py`); Redis feature store scaffold (`backend/feature_store.py`).

### Frontend note

The UI (`frontend/src/api.ts`) uses unauthenticated `fetch` for host settings and admin. If the API returns **401**, add `Authorization: Bearer <token>` from `/api/auth/login` to those calls.

"""

    write_markdown(paths, intro)

    doc = Document()
    sect = doc.sections[0]
    sect.left_margin = sect.right_margin = Pt(72)

    title = doc.add_heading("Smart Pricing (Lodging) — Complete code & explanations", 0)
    title.runs[0].font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)

    doc.add_paragraph()
    add_para(
        doc,
        f"Generated (UTC): {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M')}. "
        "Single bundle: full source listings with per-file explanations. "
        f"Total files embedded: {len(paths)}.",
        size=11,
    )
    doc.add_paragraph()
    add_para(
        doc,
        "Excluded from this bundle: node_modules, dist, __pycache__, .vite, .joblib, .zip, .csv. "
        "See also the Markdown twin: Smart_Pricing_Complete_Code_And_Explanations.md",
        size=10,
    )
    doc.add_page_break()

    # Short intro in Word (full detail in MD)
    for block in intro.split("\n"):
        if block.strip():
            add_para(doc, block.strip(), size=10)
    doc.add_page_break()

    def _part(rel: str) -> str:
        if rel.startswith("backend/"):
            return "backend"
        if rel.startswith("frontend/"):
            return "frontend"
        return "root"

    current_part: str | None = None
    for fp in paths:
        rel = fp.relative_to(ROOT).as_posix()
        part = _part(rel)
        if part != current_part:
            if part == "root":
                doc.add_heading("Part A — Project root (requirements & scripts)", level=1)
            elif part == "backend":
                doc.add_page_break()
                doc.add_heading("Part B — Backend (Python / FastAPI)", level=1)
            else:
                doc.add_page_break()
                doc.add_heading("Part C — Frontend (React / Vite / TypeScript)", level=1)
            doc.add_paragraph()
            current_part = part

        doc.add_heading(rel.replace("/", " → "), level=2)
        expl = EXPLAIN.get(rel, f"Project source file: {rel}")
        add_para(doc, expl, size=10)
        doc.add_paragraph()

        text = fp.read_text(encoding="utf-8", errors="replace")
        add_code_block_docx(doc, text)
        endp = doc.add_paragraph()
        er = endp.add_run("— end of file —")
        er.italic = True
        er.font.size = Pt(9)
        er.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
        doc.add_paragraph()

    doc.save(OUT_DOCX)
    print(f"Wrote: {OUT_DOCX}")
    print(f"Wrote: {OUT_MD}")
    print(f"Files embedded: {len(paths)}")


if __name__ == "__main__":
    main()
