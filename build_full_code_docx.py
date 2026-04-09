"""
Generate Smart_Pricing_Full_Code_Documentation.docx with complete backend + frontend source
and explanations. Excludes node_modules, dist, __pycache__.

Run: py -3.12 build_full_code_docx.py
Requires: pip install python-docx
"""
from __future__ import annotations

import os
from pathlib import Path

from docx import Document
from docx.enum.text import WD_LINE_SPACING
from docx.shared import Pt, RGBColor

ROOT = Path(__file__).resolve().parent

SKIP_DIRS = {"node_modules", "dist", "__pycache__", ".git", ".vite"}
SKIP_FILES = {".DS_Store"}

# Order matters for readability (not alphabetical)
ORDERED = [
    "backend/__init__.py",
    "backend/pricing_engine.py",
    "backend/main.py",
    "frontend/package.json",
    "frontend/package-lock.json",
    "frontend/tsconfig.json",
    "frontend/vite.config.ts",
    "frontend/index.html",
    "frontend/src/vite-env.d.ts",
    "frontend/src/api.ts",
    "frontend/src/main.tsx",
    "frontend/src/styles.css",
    "frontend/src/App.tsx",
]

EXPLAIN: dict[str, str] = {
    "backend/__init__.py": (
        "Marks the `backend` directory as a Python package so imports like `backend.main` work "
        "when you run Uvicorn from the project root."
    ),
    "backend/pricing_engine.py": (
        "Core pricing logic for lodging (nightly rates). Implements the multiplier pipeline: "
        "raw_price = base × season × demand × supply × day-of-week × lead_time × quality × goal × risk, "
        "then clamps to host min/max. Also provides confidence scores, explanation tags, blackout handling, "
        "smoothing vs the previous day, kill-switch behavior, and mock booking probability / expected revenue "
        "for the simulation API."
    ),
    "backend/main.py": (
        "FastAPI application: loads `Airbnb_Open_Data.csv`, exposes REST endpoints for listings, "
        "pricing calendar (`GET /api/pricing/{id}`), saving host settings (`POST /api/host/settings/{id}`), "
        "simulation (`POST /api/simulation/run`), admin status and kill switch. Serves the built React app "
        "from `frontend/dist` when present; otherwise returns a JSON hint at `/`. Can be run with "
        "`python main.py` inside `backend/` (fallback imports) or `uvicorn backend.main:app` from the project root."
    ),
    "frontend/package.json": (
        "NPM project metadata: React 18, Vite 5, TypeScript. Scripts `dev`, `build`, `preview` for local dev and production bundle."
    ),
    "frontend/package-lock.json": (
        "NPM lockfile: pins exact dependency versions for reproducible `npm ci` / `npm install` installs."
    ),
    "frontend/tsconfig.json": (
        "TypeScript compiler options: strict mode, ES2022, JSX react-jsx, bundler module resolution for Vite."
    ),
    "frontend/vite.config.ts": (
        "Vite configuration: React plugin, dev server on port 5173, proxy `/api` to the FastAPI backend on port 8000, "
        "build output to `dist/`."
    ),
    "frontend/index.html": (
        "HTML shell: mounts the React app on `#root`, loads DM Sans font, sets viewport and theme color."
    ),
    "frontend/src/vite-env.d.ts": (
        "TypeScript triple-slash reference so the compiler recognizes Vite-specific types and `import.meta`."
    ),
    "frontend/src/api.ts": (
        "Thin API client: `fetch` wrappers for `/api/listings`, `/api/pricing`, host settings POST, simulation, "
        "admin status, kill switch. Uses same-origin `/api` (or Vite proxy in development)."
    ),
    "frontend/src/main.tsx": (
        "React 18 entry: creates root, renders `<App />` inside `StrictMode`, imports global `styles.css`."
    ),
    "frontend/src/styles.css": (
        "Global UI styling: blue-and-white theme, layout (top bar, cards, grid), form controls, calendar cells, "
        "modal, stats, error banner. CSS variables for colors and shadows."
    ),
    "frontend/src/App.tsx": (
        "Main React UI: tabbed Host / Simulation / Admin. Host tab: listing picker, smart pricing toggles, "
        "min/max/base, preferences, locked/blackout dates, 60-day calendar with modal breakdown. Simulation tab: "
        "custom price vs probability and expected revenue. Admin tab: kill switch and audit list."
    ),
}


def collect_files() -> list[Path]:
    found: dict[str, Path] = {}
    for dirpath, dirnames, filenames in os.walk(ROOT / "backend"):
        dirnames[:] = [d for d in dirnames if d not in SKIP_DIRS]
        for fn in filenames:
            if fn in SKIP_FILES or fn.endswith(".pyc"):
                continue
            p = Path(dirpath) / fn
            rel = p.relative_to(ROOT).as_posix()
            found[rel] = p
    for dirpath, dirnames, filenames in os.walk(ROOT / "frontend"):
        dirnames[:] = [d for d in dirnames if d not in SKIP_DIRS]
        for fn in filenames:
            if fn in SKIP_FILES:
                continue
            p = Path(dirpath) / fn
            rel = p.relative_to(ROOT).as_posix()
            if "node_modules" in rel or rel.startswith("frontend/dist/"):
                continue
            found[rel] = p

    ordered_paths: list[Path] = []
    for rel in ORDERED:
        if rel in found:
            ordered_paths.append(found.pop(rel))
    # Any extra files not in ORDERED (should be none if list is complete)
    for rel in sorted(found.keys()):
        ordered_paths.append(found[rel])

    return ordered_paths


def add_para(doc: Document, text: str, *, bold: bool = False, size: int = 11) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    if bold:
        run.bold = True
    run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)


def add_code_block(doc: Document, content: str) -> None:
    """Append full source; one paragraph per line keeps Word stable for huge files."""
    if not content.endswith("\n"):
        content += "\n"
    for line in content.splitlines():
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        run = p.add_run(line)
        run.font.name = "Consolas"
        run._element.rPr.rFonts.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}ascii", "Consolas")
        run._element.rPr.rFonts.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}hAnsi", "Consolas")
        run.font.size = Pt(7.5)


def main() -> None:
    out = ROOT / "Smart_Pricing_Full_Code_Documentation.docx"

    doc = Document()
    sect = doc.sections[0]
    sect.left_margin = sect.right_margin = Pt(72)

    title = doc.add_heading("Smart Pricing (Lodging)", 0)
    title.runs[0].font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)

    doc.add_paragraph()
    add_para(
        doc,
        "Complete source code documentation — backend (Python / FastAPI) and frontend (React / TypeScript / Vite). "
        "Every file below is included in full with no omissions. Generated for client delivery.",
        size=11,
    )
    doc.add_paragraph()
    add_para(doc, "Runtime data: the API expects Airbnb_Open_Data.csv in the project root (not included in this code listing).", size=10)

    paths = collect_files()
    doc.add_page_break()

    doc.add_heading("Part A — Backend (Python)", level=1)
    doc.add_paragraph()

    backend_done = False
    for fp in paths:
        rel = fp.relative_to(ROOT).as_posix()
        if rel.startswith("frontend/") and not backend_done:
            doc.add_page_break()
            doc.add_heading("Part B — Frontend (React / Vite)", level=1)
            doc.add_paragraph()
            backend_done = True

        doc.add_heading(rel.replace("/", " → "), level=2)
        expl = EXPLAIN.get(rel, "Source file for the Smart Pricing lodging demo.")
        add_para(doc, expl, size=10)
        doc.add_paragraph()

        text = fp.read_text(encoding="utf-8", errors="replace")
        add_code_block(doc, text)
        endp = doc.add_paragraph()
        er = endp.add_run("— end of file —")
        er.italic = True
        er.font.size = Pt(9)
        er.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
        doc.add_paragraph()

    doc.save(out)
    print(f"Wrote: {out}")
    print(f"Files embedded: {len(paths)}")


if __name__ == "__main__":
    main()
